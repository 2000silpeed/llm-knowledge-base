"""Word 인제스터 (W1-05)

python-docx로 .docx 문서 파싱 → 마크다운 변환
제목 계층(H1/H2/H3) → #/##/### 변환
표 → 마크다운 테이블
각주/미주 → 문서 하단 모음
청킹: H2 단위 분할, 200토큰 overlap
출력: raw/office/{파일명}.md + raw/office/{파일명}.meta.yaml

사용 예:
    from scripts.ingest_word import ingest_word
    result = ingest_word("/path/to/file.docx")
    # {"status": "ok", "path": "raw/office/file.md", "paragraphs": 120, ...}
"""

import hashlib
import logging
import re
from datetime import datetime, timezone
from pathlib import Path

import yaml

from scripts.token_counter import estimate_tokens, load_settings
from scripts.utils import slugify as _slugify

logger = logging.getLogger(__name__)

# 헤딩 스타일명 → 마크다운 접두사
HEADING_STYLE_MAP: dict[str, str] = {
    "heading 1": "#",
    "heading 2": "##",
    "heading 3": "###",
    "heading 4": "####",
    "heading 5": "#####",
    "heading 6": "######",
}


# ──────────────────────────────────────────────
# 각주 / 미주 추출
# ──────────────────────────────────────────────

def _get_notes(doc, rel_type_suffix: str) -> dict[str, str]:
    """각주 또는 미주 텍스트를 {id: text} dict로 반환합니다.

    rel_type_suffix: "footnotes" 또는 "endnotes"
    """
    result: dict[str, str] = {}
    try:
        from docx.oxml.ns import qn

        rel_type = (
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/"
            + rel_type_suffix
        )
        notes_part = doc.part.part_related_by(rel_type)
        note_tag = qn("w:footnote") if rel_type_suffix == "footnotes" else qn("w:endnote")

        for note_elem in notes_part._element.iter(note_tag):
            note_id = note_elem.get(qn("w:id"))
            if note_id is None:
                continue
            try:
                if int(note_id) <= 0:  # -1, 0: 구분자 각주
                    continue
            except ValueError:
                continue

            text_parts: list[str] = []
            for t_elem in note_elem.iter(qn("w:t")):
                text_parts.append(t_elem.text or "")
            result[note_id] = "".join(text_parts)
    except Exception:
        pass
    return result


# ──────────────────────────────────────────────
# 단락 → 마크다운
# ──────────────────────────────────────────────

def _runs_to_text(para) -> str:
    """단락의 런을 볼드/이탤릭 처리하여 조합합니다."""
    parts: list[str] = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts)


def _find_note_refs(para, footnote_map: dict[str, str], endnote_map: dict[str, str],
                    ref_counter: dict) -> str:
    """단락에서 각주/미주 참조를 찾아 인라인 마커 문자열을 반환합니다.

    ref_counter: {"next": 1, "id_to_num": {}} — 호출 간 공유 상태
    """
    from docx.oxml.ns import qn

    markers: list[str] = []
    id_to_num: dict[str, int] = ref_counter["id_to_num"]

    for elem in para._element.iter(qn("w:footnoteReference")):
        fn_id = elem.get(qn("w:id"))
        if not fn_id or fn_id not in footnote_map:
            continue
        key = f"fn_{fn_id}"
        if key not in id_to_num:
            id_to_num[key] = ref_counter["next"]
            ref_counter["next"] += 1
        markers.append(f"[^{id_to_num[key]}]")

    for elem in para._element.iter(qn("w:endnoteReference")):
        en_id = elem.get(qn("w:id"))
        if not en_id or en_id not in endnote_map:
            continue
        key = f"en_{en_id}"
        if key not in id_to_num:
            id_to_num[key] = ref_counter["next"]
            ref_counter["next"] += 1
        markers.append(f"[^{id_to_num[key]}]")

    return "".join(markers)


def _runs_to_text_with_changes(para) -> str:
    """트랙변경 정보(삽입/삭제)를 포함해 단락 런을 변환합니다.

    - w:ins  (삽입): ++텍스트++ 로 표시
    - w:del  (삭제): ~~텍스트~~ 로 표시
    - 일반 런: 볼드/이탤릭 처리
    """
    from docx.oxml.ns import qn

    parts: list[str] = []
    body = para._element

    for child in body:
        tag = child.tag

        # 일반 런
        if tag == qn("w:r"):
            texts = [t.text or "" for t in child.iter(qn("w:t"))]
            text = "".join(texts)
            if not text:
                continue
            # 볼드/이탤릭 감지
            rpr = child.find(qn("w:rPr"))
            bold = rpr is not None and rpr.find(qn("w:b")) is not None
            italic = rpr is not None and rpr.find(qn("w:i")) is not None
            if bold and italic:
                text = f"***{text}***"
            elif bold:
                text = f"**{text}**"
            elif italic:
                text = f"*{text}*"
            parts.append(text)

        # 삽입 (tracked insertion)
        elif tag == qn("w:ins"):
            texts = [t.text or "" for t in child.iter(qn("w:t"))]
            text = "".join(texts)
            if text:
                parts.append(f"++{text}++")

        # 삭제 (tracked deletion) — w:delText 사용
        elif tag == qn("w:del"):
            texts = [t.text or "" for t in child.iter(qn("w:delText"))]
            text = "".join(texts)
            if text:
                parts.append(f"~~{text}~~")

    return "".join(parts)


def _para_to_markdown(para, footnote_map: dict[str, str], endnote_map: dict[str, str],
                      ref_counter: dict, include_tracked_changes: bool = False) -> str:
    """단락 하나를 마크다운 줄로 변환합니다. 빈 단락은 빈 문자열 반환."""
    style_name = (para.style.name or "Normal").lower() if para.style else "normal"

    # 헤딩
    if style_name in HEADING_STYLE_MAP:
        text = para.text.strip()
        if text:
            return f"{HEADING_STYLE_MAP[style_name]} {text}"
        return ""

    # 빈 단락 (트랙변경 포함 시 삭제 텍스트가 있을 수 있으므로 raw_text만으로 판단)
    raw_text = para.text.strip()
    if not raw_text and not include_tracked_changes:
        return ""

    # 런 기반 텍스트
    if include_tracked_changes:
        inline_text = _runs_to_text_with_changes(para)
    else:
        inline_text = _runs_to_text(para)
    if not inline_text.strip():
        if not raw_text:
            return ""
        inline_text = raw_text  # fallback

    # 각주/미주 참조 마커
    note_markers = _find_note_refs(para, footnote_map, endnote_map, ref_counter)
    line = inline_text + note_markers

    # 리스트 스타일 감지
    style_lower = style_name
    if "list bullet" in style_lower or "bullet" in style_lower:
        return f"- {line}"
    if "list number" in style_lower or "list paragraph" in style_lower:
        # 번호 리스트는 마크다운에서 1.로 시작하면 렌더러가 순번 처리
        return f"1. {line}"

    return line


# ──────────────────────────────────────────────
# 표 → 마크다운
# ──────────────────────────────────────────────

def _table_to_markdown(table) -> str:
    """표를 마크다운 테이블로 변환합니다."""
    rows = table.rows
    if not rows:
        return ""

    rows_md: list[str] = []
    for row_idx, row in enumerate(rows):
        cells: list[str] = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace("|", "\\|").replace("\n", " ")
            cells.append(cell_text)
        row_line = "| " + " | ".join(cells) + " |"
        rows_md.append(row_line)
        if row_idx == 0:
            separator = "| " + " | ".join(["---"] * len(cells)) + " |"
            rows_md.append(separator)

    return "\n".join(rows_md)


# ──────────────────────────────────────────────
# 문서 전체 → 마크다운 블록 목록
# ──────────────────────────────────────────────

def _doc_to_blocks(doc, footnote_map: dict[str, str], endnote_map: dict[str, str],
                   ref_counter: dict, include_tracked_changes: bool = False) -> list[str]:
    """문서 body의 단락과 표를 순서대로 마크다운 블록으로 변환합니다."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    blocks: list[str] = []
    body = doc.element.body

    for child in body:
        tag = child.tag
        if tag == qn("w:p"):
            para = Paragraph(child, doc)
            line = _para_to_markdown(
                para, footnote_map, endnote_map, ref_counter,
                include_tracked_changes=include_tracked_changes,
            )
            blocks.append(line)
        elif tag == qn("w:tbl"):
            table = Table(child, doc)
            table_md = _table_to_markdown(table)
            if table_md:
                blocks.append(table_md)

    return blocks


# ──────────────────────────────────────────────
# 각주/미주 → 문서 하단 섹션
# ──────────────────────────────────────────────

def _build_notes_section(footnote_map: dict[str, str], endnote_map: dict[str, str],
                         ref_counter: dict) -> str:
    """수집된 각주/미주를 마크다운 각주 형식으로 반환합니다."""
    id_to_num = ref_counter["id_to_num"]
    if not id_to_num:
        return ""

    # num → text 역매핑
    num_to_text: dict[int, str] = {}
    for key, num in id_to_num.items():
        if key.startswith("fn_"):
            fn_id = key[3:]
            num_to_text[num] = footnote_map.get(fn_id, "")
        elif key.startswith("en_"):
            en_id = key[3:]
            num_to_text[num] = endnote_map.get(en_id, "")

    lines: list[str] = ["---", "", "## 각주 및 미주"]
    for num in sorted(num_to_text):
        text = num_to_text[num].strip()
        if text:
            lines.append(f"[^{num}]: {text}")

    return "\n".join(lines)


# ──────────────────────────────────────────────
# H2 단위 청킹 + overlap
# ──────────────────────────────────────────────

def _split_at_h2(markdown: str) -> list[tuple[str, str]]:
    """마크다운을 H2(##) 경계에서 분할합니다.

    Returns:
        list of (section_heading, section_body)
        첫 번째 항목은 H2 이전 내용 ("__preamble__", ...)
    """
    lines = markdown.splitlines()
    sections: list[tuple[str, str]] = []
    current_heading = "__preamble__"
    current_lines: list[str] = []

    for line in lines:
        if line.startswith("## ") and not line.startswith("### "):
            sections.append((current_heading, "\n".join(current_lines)))
            current_heading = line[3:].strip()
            current_lines = [line]
        else:
            current_lines.append(line)

    sections.append((current_heading, "\n".join(current_lines)))
    return sections


def _overlap_tail(text: str, overlap_tokens: int) -> str:
    """텍스트 끝 부분에서 overlap_tokens에 해당하는 양을 반환합니다."""
    # 바이트 기반 역산: token * 4 bytes
    byte_limit = overlap_tokens * 4
    encoded = text.encode("utf-8")
    if len(encoded) <= byte_limit:
        return text
    tail_bytes = encoded[-byte_limit:]
    # UTF-8 경계 맞추기
    return tail_bytes.decode("utf-8", errors="ignore")


def _build_chunks(sections: list[tuple[str, str]], doc_name: str,
                  min_chunk_tokens: int, overlap_tokens: int) -> list[str]:
    """섹션 목록을 청크 문자열 목록으로 합칩니다.

    min_chunk_tokens보다 작은 섹션은 다음 섹션과 병합합니다.
    각 청크 앞에 청크 헤더와 overlap 블록을 삽입합니다.
    """
    # 섹션 병합: 너무 작은 섹션 합치기
    merged: list[tuple[str, str]] = []
    buffer_heading = ""
    buffer_body = ""

    for heading, body in sections:
        if not body.strip() and not buffer_body.strip():
            buffer_heading = heading
            buffer_body = body
            continue

        if buffer_body.strip():
            combined = buffer_body + "\n\n" + body
            if estimate_tokens(combined) < min_chunk_tokens and heading != "__preamble__":
                buffer_body = combined
                continue
            else:
                merged.append((buffer_heading, buffer_body))
                buffer_heading = heading
                buffer_body = body
        else:
            buffer_heading = heading
            buffer_body = body

    if buffer_body.strip() or buffer_heading:
        merged.append((buffer_heading, buffer_body))

    if not merged:
        return []

    total = len(merged)
    chunks: list[str] = []
    prev_tail = ""

    for idx, (heading, body) in enumerate(merged, start=1):
        label = heading if heading != "__preamble__" else "서문"
        header = f"[{doc_name} / {label} / {total}개 중 {idx}번째]"

        parts: list[str] = [header]
        if prev_tail:
            parts.append(f"<!-- overlap -->\n{prev_tail}\n<!-- /overlap -->")
        parts.append(body)

        chunk_text = "\n\n".join(filter(None, parts))
        chunks.append(chunk_text)
        prev_tail = _overlap_tail(body, overlap_tokens)

    return chunks


# ──────────────────────────────────────────────
# 퍼블릭 진입점
# ──────────────────────────────────────────────

def ingest_word(
    docx_path: str | Path,
    project_root: Path | str | None = None,
    settings: dict | None = None,
    include_tracked_changes: bool = False,
) -> dict:
    """Word 파일(.docx)을 인제스트합니다.

    Args:
        docx_path: Word 파일 경로
        project_root: 프로젝트 루트 경로. None이면 이 파일 기준 상위 디렉토리.
        settings: 설정 dict. None이면 settings.yaml 자동 로드.
        include_tracked_changes: True이면 삽입(++text++)/삭제(~~text~~) 표시 포함.

    Returns:
        {
            "status": "ok" | "error",
            "path": str,          # 저장된 .md 파일 경로
            "meta_path": str,     # 저장된 .meta.yaml 파일 경로
            "title": str,
            "paragraphs": int,
            "tables": int,
            "footnotes": int,
            "endnotes": int,
            "chunk_count": int,
            "token_count": int,
            "message": str,       # 오류 시 메시지
        }
    """
    try:
        from docx import Document
    except ImportError:
        return {"status": "error", "message": "python-docx가 설치되지 않았습니다. `uv add python-docx`"}

    if project_root is None:
        project_root = Path(__file__).parent.parent
    project_root = Path(project_root)
    docx_path = Path(docx_path)

    if not docx_path.exists():
        return {"status": "error", "message": f"파일 없음: {docx_path}"}
    if docx_path.suffix.lower() not in {".docx", ".doc"}:
        return {"status": "error", "message": f"Word 파일이 아닙니다: {docx_path}"}
    if docx_path.suffix.lower() == ".doc":
        return {
            "status": "error",
            "message": ".doc (구형 바이너리) 포맷은 python-docx 미지원. .docx로 변환 후 시도하세요.",
        }

    if settings is None:
        settings = load_settings()

    office_dir = project_root / settings["paths"]["raw"] / "office"
    office_dir.mkdir(parents=True, exist_ok=True)

    overlap_tokens = settings.get("chunking", {}).get("overlap_tokens", 200)
    min_chunk_tokens = settings.get("chunking", {}).get("min_chunk_tokens", 500)

    logger.info("Word 문서 로드 중: %s", docx_path)

    try:
        doc = Document(str(docx_path))
    except Exception as exc:
        return {"status": "error", "message": f"Word 열기 실패: {exc}"}

    collected_at = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    # 문서 제목: 코어 속성 또는 첫 번째 H1 또는 파일명
    title = docx_path.stem
    try:
        core_title = doc.core_properties.title
        if core_title and core_title.strip():
            title = core_title.strip()
    except Exception:
        pass

    if title == docx_path.stem:
        # 코어 속성 없으면 첫 번째 Heading 1 시도
        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower() if para.style else ""
            if style_name == "heading 1" and para.text.strip():
                title = para.text.strip()
                break

    # 각주 / 미주 추출
    footnote_map = _get_notes(doc, "footnotes")
    endnote_map = _get_notes(doc, "endnotes")

    ref_counter: dict = {"next": 1, "id_to_num": {}}

    # 문서 블록 변환
    blocks = _doc_to_blocks(
        doc, footnote_map, endnote_map, ref_counter,
        include_tracked_changes=include_tracked_changes,
    )

    # 연속된 빈 줄 정리 (최대 1개 빈 줄)
    body_lines: list[str] = []
    prev_empty = False
    for line in blocks:
        is_empty = not line.strip()
        if is_empty and prev_empty:
            continue
        body_lines.append(line)
        prev_empty = is_empty

    body_text = "\n".join(body_lines).strip()

    # 각주/미주 하단 섹션
    notes_section = _build_notes_section(footnote_map, endnote_map, ref_counter)

    full_text = body_text
    if notes_section:
        full_text = body_text + "\n\n" + notes_section

    # 통계
    para_count = len([p for p in doc.paragraphs if p.text.strip()])
    table_count = len(doc.tables)
    fn_count = len([k for k in ref_counter["id_to_num"] if k.startswith("fn_")])
    en_count = len([k for k in ref_counter["id_to_num"] if k.startswith("en_")])

    # H2 단위 청킹
    doc_name = title[:40]  # 청크 헤더용 짧은 이름
    sections = _split_at_h2(full_text)
    chunk_texts = _build_chunks(sections, doc_name, min_chunk_tokens, overlap_tokens)
    chunk_count = len(chunk_texts)

    if chunk_count > 1:
        body_with_chunks = "\n\n---\n\n".join(chunk_texts)
    else:
        body_with_chunks = chunk_texts[0] if chunk_texts else full_text

    token_count = estimate_tokens(body_with_chunks)

    # frontmatter
    frontmatter: dict = {
        "title": title,
        "source_file": str(docx_path),
        "collected_at": collected_at,
        "paragraphs": para_count,
        "tables": table_count,
        "footnotes": fn_count,
        "endnotes": en_count,
        "chunk_count": chunk_count,
        "token_count": token_count,
    }
    if include_tracked_changes:
        frontmatter["tracked_changes"] = True
    fm_str = yaml.dump(frontmatter, allow_unicode=True, default_flow_style=False, sort_keys=False)
    document = f"---\n{fm_str}---\n\n# {title}\n\n{body_with_chunks}\n"

    # 파일명 결정 (충돌 방지)
    stem = _slugify(docx_path.stem) or "document"
    filename = f"{stem}.md"
    dest = office_dir / filename
    if dest.exists():
        file_hash = hashlib.md5(docx_path.stem.encode()).hexdigest()[:6]
        filename = f"{stem}_{file_hash}.md"
        dest = office_dir / filename

    dest.write_text(document, encoding="utf-8")
    rel_path = str(dest.relative_to(project_root))

    # .meta.yaml
    meta: dict = {
        "source_file": str(docx_path),
        "collected_at": collected_at,
        "title": title,
        "paragraph_count": para_count,
        "table_count": table_count,
        "footnote_count": fn_count,
        "endnote_count": en_count,
        "chunk_count": chunk_count,
        "overlap_tokens": overlap_tokens,
        "token_count": token_count,
        "sections": [
            {"heading": h, "tokens": estimate_tokens(b)}
            for h, b in sections
            if b.strip()
        ],
    }
    meta_filename = dest.stem + ".meta.yaml"
    meta_dest = office_dir / meta_filename
    meta_dest.write_text(
        yaml.dump(meta, allow_unicode=True, default_flow_style=False, sort_keys=False),
        encoding="utf-8",
    )
    meta_rel_path = str(meta_dest.relative_to(project_root))

    logger.info(
        "저장 완료: %s (단락: %d, 표: %d, 각주: %d, 청크: %d, 토큰: %d)",
        rel_path, para_count, table_count, fn_count, chunk_count, token_count,
    )

    return {
        "status": "ok",
        "path": rel_path,
        "meta_path": meta_rel_path,
        "title": title,
        "paragraphs": para_count,
        "tables": table_count,
        "footnotes": fn_count,
        "endnotes": en_count,
        "chunk_count": chunk_count,
        "token_count": token_count,
    }
